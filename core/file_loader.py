"""Extract plain text from uploaded files for any demo.

Supports .txt, .md, .pdf, .docx. Accepts file-like objects as produced by
Streamlit's `st.file_uploader`.
"""

from __future__ import annotations

import io
from typing import IO


def load_text(filename: str, data: bytes) -> str:
    """Return the text contents of a supported file.

    Args:
        filename: Original filename (used only for extension detection).
        data: Raw bytes of the file.
    """
    name = filename.lower()

    if name.endswith((".txt", ".md")):
        return data.decode("utf-8", errors="replace")

    if name.endswith(".pdf"):
        return _load_pdf(io.BytesIO(data))

    if name.endswith(".docx"):
        return _load_docx(io.BytesIO(data))

    raise ValueError(
        f"Unsupported file type: {filename}. Supported: .txt, .md, .pdf, .docx"
    )


def _load_pdf(stream: IO[bytes]) -> str:
    from pypdf import PdfReader

    reader = PdfReader(stream)
    return "\n\n".join((page.extract_text() or "") for page in reader.pages).strip()


def _load_docx(stream: IO[bytes]) -> str:
    import docx

    document = docx.Document(stream)
    return "\n".join(p.text for p in document.paragraphs).strip()


def docx_bytes_from_text(text: str) -> bytes:
    """Produce a minimal .docx file from plain text (paragraphs split on blank lines)."""
    import docx

    doc = docx.Document()
    for block in text.split("\n\n"):
        doc.add_paragraph(block)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
